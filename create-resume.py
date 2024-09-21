from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_style(
    run,
    font_name="Calibri",
    font_size=11,
    bold=False,
    italic=False,
    color=RGBColor(0, 0, 0),
):
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic
    font.color.rgb = color


def add_section_title(doc, title, level=1):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title.upper())
    set_run_style(run, font_size=13, bold=True, color=RGBColor(0, 51, 102))
    paragraph.style = f"Heading {level}"
    paragraph.space_before = Pt(12)
    paragraph.space_after = Pt(4)  # Reduced space after section titles
    return paragraph


def add_content(doc, content, style="Normal"):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(content)
    set_run_style(run)
    return paragraph


def add_bullet_point(doc, content):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(content)
    set_run_style(run)
    paragraph.space_after = Pt(2)  # Reduced space between bullet points
    return paragraph


def create_table(doc, rows, cols):
    table = doc.add_table(rows=rows, cols=cols)
    table.allow_autofit = False
    return table


def set_column_width(column, width):
    for cell in column.cells:
        cell.width = width


def add_horizontal_line(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    r = run._r
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    border.append(bottom)
    r.append(border)
    return paragraph


def create_resume():
    doc = Document()

    # Document styling
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Custom styles
    styles = doc.styles
    style = styles.add_style("Name", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Calibri"
    style.font.size = Pt(24)
    style.font.color.rgb = RGBColor(0, 51, 102)
    style.paragraph_format.space_after = Pt(0)

    style = styles.add_style("ContactInfo", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0, 51, 102)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(6)

    # Header
    name = doc.add_paragraph("Umer Farooq", style="Name")
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact Information
    contact_info = doc.add_paragraph(style="ContactInfo")
    contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info.add_run(
        "Email: umerfarooq.dev@gmail.com | Phone: +923404744007 | LinkedIn: [Your LinkedIn URL]"
    )

    # Horizontal line
    add_horizontal_line(doc)

    # Main content table
    table = create_table(doc, rows=1, cols=2)
    set_column_width(table.columns[0], Inches(5.5))  # Main content width
    set_column_width(table.columns[1], Inches(2.0))  # Sidebar width

    main_cell = table.cell(0, 0)
    sidebar_cell = table.cell(0, 1)

    # Main content
    main_content(main_cell)

    # Sidebar
    sidebar_content(sidebar_cell)

    # Save the document
    doc.save("Umer_Farooq_Resume.docx")
    print("Resume has been created successfully!")


def main_content(cell):
    doc = cell.add_paragraph().part.document

    # Summary
    add_section_title(doc, "Professional Summary")
    summary = add_content(
        doc,
        "Senior Software Engineer & DevOps Specialist with 6+ years of experience in software engineering, DevOps, and infrastructure management. Expertise in designing scalable microservices architectures, full-stack development (MERN stack, Python, Golang), and implementing CI/CD practices. Proficient in various databases and cloud technologies, with a focus on AWS services.",
    )
    summary.space_after = Pt(6)  # Add some space after the summary

    # Work Experience
    add_section_title(doc, "Professional Experience")

    # Two-column layout for work experience
    exp_table = create_table(doc, rows=3, cols=2)
    set_column_width(exp_table.columns[0], Inches(1.5))  # Date column
    set_column_width(exp_table.columns[1], Inches(4.0))  # Description column

    # Kaleidoscope experience (enhanced visibility)
    exp_table.cell(0, 0).text = "Jan 2023 - Present"
    exp_cell = exp_table.cell(0, 1)
    company_name = exp_cell.paragraphs[0].add_run(
        "Senior Software Engineer / DevOps, Kaleidoscope"
    )
    set_run_style(
        company_name, bold=True, font_size=12
    )  # Make company name bold and slightly larger
    for bullet in [
        "Led development of HypergraphTM platform, providing end-to-end observability across operational sources",
        "Implemented crawlers for AWS services (S3, IAM) metadata retrieval and analyzers for security insights",
        "Utilized AWS Go SDK, Redux, and Cytoscape for complex data visualization and state management",
        "Implemented ETL job stages using Golang's goflow package for efficient long-running tasks",
        "Integrated Keycloak for SSO and authentication, enhancing security and user experience",
        "Managed CI/CD pipelines using GitHub Actions and Jenkins for multiple repositories",
        "Orchestrated microservices deployment on Amazon EKS, utilizing GitHub Container Registry",
    ]:
        add_bullet_point(exp_cell, bullet)

    # TheMoonShotFactory experience
    exp_table.cell(1, 0).text = "Mar 2020 - Jan 2023"
    exp_cell = exp_table.cell(1, 1)
    company_name = exp_cell.paragraphs[0].add_run(
        "Senior Software Engineer, TheMoonShotFactory (CrownBio Project)"
    )
    set_run_style(company_name, bold=True)
    for bullet in [
        "Developed and maintained Crownlink Saturn Application for laboratory test data management in the biotech domain",
        "Implemented AWS services including DynamoDB, AppSync, Lambda, Cognito, and CloudFront",
        "Created data pipelines using AWS Glue, S3, SQS, and SNS for analytics and event processing",
        "Utilized AWS DAG for managing ETL jobs and visual representation of data pipelines",
        "Implemented event-based invocation on Lambda functions for real-time processing",
        "Developed a Django application integrated with Amazon SQS and AWS Redshift for real-time message processing",
    ]:
        add_bullet_point(exp_cell, bullet)

    # Veeve.io experience
    exp_table.cell(2, 0).text = "Sep 2018 - Mar 2020"
    exp_cell = exp_table.cell(2, 1)
    company_name = exp_cell.paragraphs[0].add_run("Full Stack MERN Developer, Veeve.io")
    set_run_style(company_name, bold=True)
    for bullet in [
        "Developed Veeve Smart Cart technology for grocery retailers, enabling scan-and-go shopping experiences",
        "Implemented weight sensor and product scanner integration for accurate cart management",
        "Created a product recommendation system based on current store inventory and user selections",
        "Integrated Stripe payment gateway for secure transactions",
        "Utilized Redis for caching and MongoDB Change Streams for real-time updates and cache invalidation",
        "Implemented comprehensive unit tests covering positive and negative edge cases",
    ]:
        add_bullet_point(exp_cell, bullet)


def sidebar_content(cell):
    doc = cell.add_paragraph().part.document

    # Skills (more compact layout)
    add_section_title(doc, "Skills")
    skills = [
        "JavaScript / TypeScript",
        "React.js / Redux",
        "Node.js / Express.js",
        "Python / Django",
        "Golang",
        "AWS Services",
        "Docker / Kubernetes",
        "CI/CD (Jenkins, GitHub Actions)",
        "MongoDB / PostgreSQL",
        "GraphQL / REST APIs",
        "Microservices Architecture",
        "DevOps Practices",
    ]
    skills_table = create_table(doc, rows=4, cols=3)
    for i, skill in enumerate(skills):
        skills_table.cell(i // 3, i % 3).text = skill

    # Add some space after the skills table
    doc.add_paragraph().space_after = Pt(6)

    # Education
    add_section_title(doc, "Education")
    add_content(doc, "Bachelor of Information Technology")
    add_content(doc, "Government College University Faisalabad, 2016 - 2020")
    add_bullet_point(doc, "Graduated in top 10% of class")
    add_bullet_point(doc, "EU Blue Card eligible degree")

    # Certifications
    add_section_title(doc, "Certifications")
    add_bullet_point(doc, "AWS Certified Solutions Architect - Associate (2023)")
    add_bullet_point(doc, "AWS Certified Developer - Associate (2023)")
    add_bullet_point(doc, "AWS Certified DevOps Engineer - Professional (2023)")

    # Personal Projects
    add_section_title(doc, "Personal Projects")
    add_content(doc, "Animated Backgrounds for React")
    add_bullet_point(doc, "React package for customizable animated backgrounds")
    add_bullet_point(doc, "Published on npm with documentation")
    add_content(doc, "FoodShare Web Application")
    add_bullet_point(doc, "Full-stack app for sharing food with people in need")
    add_bullet_point(doc, "Used React, Node.js, and Google Maps API")
    # Deployment Projects
    add_section_title(doc, "Deployment Projects")
    add_content(doc, "Deployed Applications on Private VPS")
    add_bullet_point(
        doc, "Deployed Python applications from scratch using Nginx and Caddy"
    )
    add_bullet_point(doc, "Deployed Next.js, React, and Express.js applications on VPS")
    add_bullet_point(
        doc, "Configured and managed server environments for optimal performance"
    )
    add_bullet_point(doc, "Ensured secure and scalable deployment processes")


if __name__ == "__main__":
    create_resume()
